# -*- coding: utf-8 -*-
"""Помощники сборки .docx по стандарту doc-style (профиль neutral)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import copy

NBSP = " "

# Токены профиля neutral
T = {
    "heading": "1A1A1A", "body": "333333", "muted": "808080",
    "line": "D9D9D9", "accent": "1F3A5F", "accent2": "2C3E50",
    "section_light": "D6E0EC", "callout": "F5F5F5", "warn": "C77B30",
    "font": "Roboto",
}

def rgb(hexstr):
    return RGBColor.from_string(hexstr)

def set_font(run, size=11, bold=False, italic=False, color=T["body"], font=None):
    run.font.name = font or T["font"]
    r = run._element.rPr
    rF = r.find(qn('w:rFonts'))
    if rF is None:
        rF = OxmlElement('w:rFonts'); r.append(rF)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rF.set(qn(a), font or T["font"])
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)

def para(doc, text=None, size=11, bold=False, italic=False, color=T["body"],
         align="justify", before=0, after=6, style=None, keep_next=False):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15
    pf.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    if keep_next:
        kn = OxmlElement('w:keepNext'); p._p.get_or_add_pPr().append(kn)
    if text is not None:
        # **bold** inline markup
        parts = text.split("**")
        for i, chunk in enumerate(parts):
            if chunk == "":
                continue
            r = p.add_run(chunk)
            set_font(r, size=size, bold=bold or (i % 2 == 1), italic=italic, color=color)
    return p

# ---------- нумерация (штатный механизм Word) ----------

_num_counter = [100]

def _numbering_part(doc):
    return doc.part.numbering_part.element

def add_multilevel_heading_numbering(doc):
    """Многоуровневая нумерация 1. / 1.1. / 1.1.1. для заголовков."""
    numbering = _numbering_part(doc)
    absId = 90
    absNum = OxmlElement('w:abstractNum')
    absNum.set(qn('w:abstractNumId'), str(absId))
    for lvl_i, fmt in enumerate(["%1.", "%1.%2.", "%1.%2.%3."]):
        lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), str(lvl_i))
        start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
        nf = OxmlElement('w:numFmt'); nf.set(qn('w:val'), 'decimal'); lvl.append(nf)
        lt = OxmlElement('w:lvlText'); lt.set(qn('w:val'), fmt); lvl.append(lt)
        lj = OxmlElement('w:lvlJc'); lj.set(qn('w:val'), 'left'); lvl.append(lj)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '0'); ind.set(qn('w:firstLine'), '0')
        pPr.append(ind); lvl.append(pPr)
        suff = OxmlElement('w:suff'); suff.set(qn('w:val'), 'space'); lvl.append(suff)
        absNum.append(lvl)
    numbering.insert(0, absNum)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(absId))
    a = OxmlElement('w:abstractNumId'); a.set(qn('w:val'), str(absId)); num.append(a)
    numbering.append(num)
    return absId

def add_list_numbering(doc, kind="bullet", level1_char="—"):
    """Отдельный numbering для каждого независимого списка. kind: bullet|decimal."""
    numbering = _numbering_part(doc)
    _num_counter[0] += 1
    absId = _num_counter[0]
    absNum = OxmlElement('w:abstractNum')
    absNum.set(qn('w:abstractNumId'), str(absId))
    levels = [(0, 360, 720), (1, 720, 1080)]
    chars = [level1_char, "•"]
    for lvl_i, hang_pos, text_pos in levels:
        lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), str(lvl_i))
        start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
        nf = OxmlElement('w:numFmt')
        nf.set(qn('w:val'), 'bullet' if kind == 'bullet' else 'decimal')
        lvl.append(nf)
        lt = OxmlElement('w:lvlText')
        lt.set(qn('w:val'), chars[lvl_i] if kind == 'bullet' else ("%1." if lvl_i == 0 else "%1.%2."))
        lvl.append(lt)
        lj = OxmlElement('w:lvlJc'); lj.set(qn('w:val'), 'left'); lvl.append(lj)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(text_pos)); ind.set(qn('w:hanging'), '360')
        pPr.append(ind); lvl.append(pPr)
        if kind == 'bullet':
            rPr = OxmlElement('w:rPr')
            rF = OxmlElement('w:rFonts')
            for a in ('w:ascii', 'w:hAnsi'):
                rF.set(qn(a), T["font"])
            rPr.append(rF); lvl.append(rPr)
        absNum.append(lvl)
    numbering.append(absNum)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(absId))
    a = OxmlElement('w:abstractNumId'); a.set(qn('w:val'), str(absId)); num.append(a)
    numbering.append(num)
    return absId

def list_item(doc, text, numId, level=0, after=3, size=11):
    p = para(doc, text, size=size, align="justify", after=after)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), str(level)); numPr.append(ilvl)
    numEl = OxmlElement('w:numId'); numEl.set(qn('w:val'), str(numId)); numPr.append(numEl)
    pPr.append(numPr)
    return p

def heading(doc, text, level=1, numId=None):
    spacing = {1: (18, 10), 2: (14, 8), 3: (10, 6)}[level]
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(spacing[0]); pf.space_after = Pt(spacing[1])
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.style = doc.styles[f'Heading {level}']
    kn = OxmlElement('w:keepNext'); p._p.get_or_add_pPr().append(kn)
    if numId:
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), str(level - 1)); numPr.append(ilvl)
        numEl = OxmlElement('w:numId'); numEl.set(qn('w:val'), str(numId)); numPr.append(numEl)
        pPr.append(numPr)
    r = p.add_run(text)
    set_font(r, size=12, bold=True, color=T["heading"])
    return p

# ---------- таблицы ----------

def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)

def _set_borders(el_pr, spec):
    """spec: dict side -> (sz_eighths, color) or None"""
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        v = spec.get(side)
        if v is None:
            b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0')
        else:
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(v[0]))
            b.set(qn('w:color'), v[1])
        borders.append(b)
    el_pr.append(borders)

def _cell_border(cell, side, sz, color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
    b = tcB.find(qn(f'w:{side}'))
    if b is None:
        b = OxmlElement(f'w:{side}'); tcB.append(b)
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz)); b.set(qn('w:color'), color)

def cell_text(cell, text, size=10, bold=False, color=T["body"], align="left", after=0):
    cell.text = ""
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(after); pf.line_spacing = 1.1
    pf.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    parts = text.split("**")
    for i, chunk in enumerate(parts):
        if chunk == "":
            continue
        r = p.add_run(chunk)
        set_font(r, size=size, bold=bold or (i % 2 == 1), color=color)

def framed_table(doc, headers, rows, widths_cm, num_cols_right=()):
    """Текстовая таблица с рамками, чередованием строк."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    _set_borders(tbl._tbl.tblPr, {s: (8, T["line"]) for s in
                                  ('top', 'left', 'bottom', 'right', 'insideH', 'insideV')})
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.width = Cm(widths_cm[j])
        _set_cell_bg(c, T["section_light"]); _set_cell_margins(c)
        cell_text(c, h, bold=True, color=T["heading"])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.rows[i + 1].cells[j]
            c.width = Cm(widths_cm[j])
            _set_cell_margins(c)
            if i % 2 == 1:
                _set_cell_bg(c, T["callout"])
            cell_text(c, val, align="right" if j in num_cols_right else "left")
    return tbl

def minimal_table(doc, headers, rows, widths_cm, num_cols_right=()):
    """Числовая таблица без вертикальных линий."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    _set_borders(tbl._tbl.tblPr, {})  # все none, границы зададим по ячейкам
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.width = Cm(widths_cm[j])
        _set_cell_margins(c)
        _cell_border(c, 'top', 16, T["heading"])
        _cell_border(c, 'bottom', 16, T["heading"])
        cell_text(c, h, bold=True, color=T["heading"],
                  align="right" if j in num_cols_right else "left")
    n = len(rows)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.rows[i + 1].cells[j]
            c.width = Cm(widths_cm[j])
            _set_cell_margins(c)
            if i < n - 1:
                _cell_border(c, 'bottom', 8, T["line"])
            else:
                _cell_border(c, 'bottom', 16, T["heading"])
            cell_text(c, val, align="right" if j in num_cols_right else "left")
    return tbl

def caption(doc, text, before=0, after=4, keep_next=True):
    return para(doc, text, size=9, italic=True, color=T["muted"], align="left",
                before=before, after=after, keep_next=keep_next)

def callout(doc, title, body_paras, accent=None):
    """Одноячеечная таблица-выноска."""
    accent = accent or T["accent"]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    _set_borders(tbl._tbl.tblPr, {})
    row = tbl.rows[0]
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit'); trPr.append(cs)
    c = row.cells[0]
    c.width = Cm(16.0)
    _set_cell_bg(c, T["callout"])
    _set_cell_margins(c, top=140, bottom=140, left=200, right=200)
    c.text = ""
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3); p0.paragraph_format.space_before = Pt(0)
    r = p0.add_run(title)
    set_font(r, size=9, bold=True, color=accent)
    for i, btext in enumerate(body_paras):
        p = c.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0 if i == len(body_paras) - 1 else 4)
        pf.line_spacing = 1.15
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parts = btext.split("**")
        for k, chunk in enumerate(parts):
            if chunk == "":
                continue
            rr = p.add_run(chunk)
            set_font(rr, size=11, bold=(k % 2 == 1))
    return tbl

def kpi_plate(doc, number, unit, label):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    _set_borders(tbl._tbl.tblPr, {})
    row = tbl.rows[0]
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit'); trPr.append(cs)
    c = row.cells[0]
    c.width = Cm(16.0)
    _set_cell_bg(c, T["callout"])
    _set_cell_margins(c, top=160, bottom=160, left=220, right=220)
    c.text = ""
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    r = p0.add_run(number)
    set_font(r, size=32, bold=True, color=T["accent"])
    if unit:
        r2 = p0.add_run(" " + unit)
        set_font(r2, size=16, bold=True, color=T["accent"])
    p1 = c.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    r3 = p1.add_run(label)
    set_font(r3, size=11, color=T["muted"])
    return tbl

# ---------- документ ----------

def new_document(header_left, header_right):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(11906); sec.page_height = Twips(16838)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)

    st = doc.styles['Normal']
    st.font.name = T["font"]; st.font.size = Pt(11)
    st.font.color.rgb = rgb(T["body"])
    rPr = st.element.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rF.set(qn(a), T["font"])
    rPr.append(rF)

    for lvl in (1, 2, 3):
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = T["font"]; hs.font.size = Pt(12); hs.font.bold = True
        hs.font.color.rgb = rgb(T["heading"])
        hrPr = hs.element.get_or_add_rPr()
        hrF = OxmlElement('w:rFonts')
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
            hrF.set(qn(a), T["font"])
        hrPr.append(hrF)

    # верхний колонтитул
    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT)
    r1 = hp.add_run(header_left); set_font(r1, size=9, bold=True, color=T["muted"])
    r2 = hp.add_run("\t" + header_right); set_font(r2, size=9, color=T["muted"])
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm = OxmlElement('w:bottom')
    btm.set(qn('w:val'), 'single'); btm.set(qn('w:sz'), '8'); btm.set(qn('w:color'), T["line"])
    pBdr.append(btm); pPr.append(pBdr)

    # нижний колонтитул: Стр. N справа, линейка сверху
    ftr = sec.footer
    fp = ftr.paragraphs[0]
    fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(2)
    rt = fp.add_run("Стр. "); set_font(rt, size=9, color=T["muted"])
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    rf = fp.add_run(); set_font(rf, size=9, color=T["muted"])
    rf._element.append(fld1)
    rf2 = fp.add_run(); set_font(rf2, size=9, color=T["muted"])
    rf2._element.append(instr)
    rf3 = fp.add_run(); set_font(rf3, size=9, color=T["muted"])
    rf3._element.append(fld2)
    fpPr = fp._p.get_or_add_pPr()
    fBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '8'); top.set(qn('w:color'), T["line"])
    fBdr.append(top); fpPr.append(fBdr)
    return doc

def doc_title(doc, title, subtitle=None, date_str=None):
    p = para(doc, None, align="left", before=0, after=2)
    r = p.add_run(title)
    set_font(r, size=14, bold=True, color=T["accent2"])
    if subtitle:
        para(doc, subtitle, size=11, color=T["muted"], align="left", after=2)
    if date_str:
        para(doc, date_str, size=9, italic=True, color=T["muted"], align="left", after=10)
